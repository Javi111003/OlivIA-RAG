# src/data_preparation/document_loader.py

import os
import re
from typing import Union, List
import io
from PIL import Image
import pytesseract as ptsract

from pypdf import PdfReader
from pypdf.errors import PdfReadError

from docx import Document
from docx.opc.exceptions import OpcError

from bs4 import BeautifulSoup
import lxml

import fitz  # PyMuPDF es importado como fitz

import src.data_preparation.config as config 

class DocumentLoader:
    """
    Clase para cargar y extraer texto de diferentes tipos de documentos (PDF, DOCX, HTML, Markdown).
    Soporta OCR para PDFs escaneados si Tesseract está configurado.
    """
    def __init__(self, tesseract_cmd_path: str = None):
        """
        Inicializa el cargador de documentos.
        :param tesseract_cmd_path: Ruta al ejecutable de Tesseract OCR si se usa para PDFs/Imágenes.
                                   Si es None, se asume que Tesseract está en el PATH del sistema.
        """
        if ptsract and tesseract_cmd_path:
            ptsract.pytesseract.tesseract_cmd = tesseract_cmd_path
        elif ptsract and not tesseract_cmd_path:
            try:
                ptsract.pytesseract.get_tesseract_version()
            except ptsract.TesseractNotFoundError:
                print("Advertencia: Tesseract OCR no se encontró en el PATH. OCR no funcionará a menos que la ruta se especifique en config.py.")
                pytsract = None 
        self.tesseract_enabled = ptsract is not None

    def _extract_text_from_pdf_page_with_ocr(self, page_image: Image.Image) -> str:
        """
        Extrae texto de una imagen de página usando Tesseract OCR.
        """
        if not self.tesseract_enabled:
            return ""
        try:
            # lang='spa' para español, puedes añadir 'eng' si hay mezcla de idiomas
            return ptsract.image_to_string(page_image, lang='spa')
        except Exception as e:
            print(f"Error durante OCR en la imagen: {e}")
            return ""

    def load_pdf(self, file_path: str) -> str:
        """
        Carga un archivo PDF y extrae su texto.
        Prioriza la extracción de texto nativo. Si no hay texto, intenta OCR si está disponible.
        Requiere PyPDF2 para texto nativo y PyMuPDF (fitz) + Tesseract para OCR robusto.
        """
        full_text = []
        try:
            # Intento 1: Extracción de texto nativo con pypdf
            if PdfReader:
                reader = PdfReader(file_path)
                for page_num in range(len(reader.pages)):
                    page_text = reader.pages[page_num].extract_text()
                    if page_text:
                        full_text.append(page_text)
                    else:
                        # Si no hay texto nativo, es probable que sea una imagen (escaneado)
                        if self.tesseract_enabled and fitz:
                            print(f"Página {page_num+1} de '{os.path.basename(file_path)}' parece escaneada. Intentando OCR...")
                            try:
                                doc = fitz.open(file_path)
                                page = doc[page_num]
                                pix = page.get_pixmap()
                                img = Image.open(io.BytesIO(pix.pil_tobytes(format="PNG"))) # Convertir a PIL Image
                                ocr_text = self._extract_text_from_pdf_page_with_ocr(img)
                                full_text.append(ocr_text)
                            except Exception as e_ocr:
                                print(f"Error durante OCR en página {page_num+1}: {e_ocr}")
                                full_text.append("") # Añadir vacío si falla el OCR
                        else:
                            full_text.append("") # Si no hay OCR, añadir vacío
            else:
                print("pypdf no está disponible para carga de PDF.")
                return ""

            # Post-procesamiento básico para eliminar posibles duplicaciones de texto en PDFs problemáticos
            # Esto es un patrón común donde el texto nativo y el OCR pueden superponerse o ser pobres.
            combined_text = "\n".join(full_text)
            # Elimina múltiples saltos de línea y espacios en blanco introducidos por la concatenación o OCR
            combined_text = re.sub(r'\s*\n\s*', '\n', combined_text).strip()
            combined_text = re.sub(r'[ \t]+', ' ', combined_text)
            return combined_text

        except PdfReadError:
            print(f"Error: El archivo PDF '{file_path}' está corrupto o no es un PDF válido.")
            return ""
        except Exception as e:
            print(f"Error inesperado al cargar PDF '{file_path}': {e}")
            return ""


    def load_docx(self, file_path: str) -> str:
        """
        Carga un archivo DOCX y extrae su texto.
        """
        if not Document:
            print("python-docx no está disponible para carga de DOCX.")
            return ""
        try:
            document = Document(file_path)
            full_text = []
            for paragraph in document.paragraphs:
                full_text.append(paragraph.text)
            return "\n".join(full_text)
        except OpcError:
            print(f"Error: El archivo DOCX '{file_path}' está corrupto o no es un DOCX válido.")
            return ""
        except Exception as e:
            print(f"Error inesperado al cargar DOCX '{file_path}': {e}")
            return ""

    def load_html(self, file_path: str) -> str:
        """
        Carga un archivo HTML y extrae el texto principal, eliminando etiquetas HTML.
        """
        if not BeautifulSoup:
            print("beautifulsoup4 no está disponible para carga de HTML.")
            return ""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()
            soup = BeautifulSoup(html_content, 'lxml') # Usamos 'lxml' como parser
            return soup.get_text(separator='\n') # Extrae texto y separa con saltos de línea
        except FileNotFoundError:
            print(f"Error: Archivo HTML no encontrado: '{file_path}'")
            return ""
        except Exception as e:
            print(f"Error inesperado al cargar HTML '{file_path}': {e}")
            return ""

    def load_markdown(self, file_path: str) -> str:
        """
        Carga un archivo Markdown y extrae su texto plano.
        """
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except FileNotFoundError:
            print(f"Error: Archivo Markdown no encontrado: '{file_path}'")
            return ""
        except Exception as e:
            print(f"Error inesperado al cargar Markdown '{file_path}': {e}")
            return ""

    def load_document(self, file_path: str) -> Union[str, None]:
        """
        Detecta el tipo de archivo basándose en la extensión y llama al cargador adecuado.
        """
        if not os.path.exists(file_path):
            print(f"Error: El archivo no existe: '{file_path}'")
            return None

        _, ext = os.path.splitext(file_path)
        ext = ext.lower()

        if ext == '.pdf':
            return self.load_pdf(file_path)
        elif ext == '.docx':
            return self.load_docx(file_path)
        elif ext == '.html' or ext == '.htm':
            return self.load_html(file_path)
        elif ext == '.md':
            return self.load_markdown(file_path)
        elif ext == '.txt': # Añadido soporte para archivos de texto plano
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            except Exception as e:
                print(f"Error al cargar TXT {file_path}: {e}")
                return ""
        else:
            print(f"Tipo de archivo no soportado: '{ext}' para '{file_path}'")
            return None

# Ejemplo de uso (esto es solo para demostración, la lógica principal iría en pipeline.py)
if __name__ == "__main__":
    # Asegúrate de que tienes un archivo config.py en la raíz del proyecto
    # y que RUTA_TESSERACT_CMD está configurado si necesitas OCR.

    # Ejemplo de instanciación del loader
    loader = DocumentLoader(tesseract_cmd_path=config.RUTA_TESSERACT_CMD)

    # --- Creación de archivos de prueba (NO HAGAS ESTO EN PRODUCCIÓN) ---
    # Esto es solo para que puedas probar el loader si no tienes archivos reales
    # Asegúrate de que los directorios de data/raw existen
    os.makedirs(os.path.join(config.RAW_DATA_DIR, 'test_subject'), exist_ok=True)

    ## DOCX de prueba
    try:
        if Document:
            doc = Document()
            doc.add_paragraph("Este es un párrafo de prueba en un documento DOCX.")
            test_docx_path = os.path.join(config.RAW_DATA_DIR, 'test_subject', 'test_doc.docx')
            doc.save(test_docx_path)
            print(f"Creado DOCX de prueba: {test_docx_path}")
    except Exception as e:
        print(f"No se pudo crear DOCX de prueba (python-docx no disponible?): {e}")

    # HTML de prueba
    test_html_path = os.path.join(config.RAW_DATA_DIR, 'test_subject', 'test_page.html')
    with open(test_html_path, 'w', encoding='utf-8') as f:
        f.write("<html><body><h1>Título</h1><p>Este es <b>un párrafo</b> en HTML.</p></body></html>")
    print(f"Creado HTML de prueba: {test_html_path}")

    # TXT de prueba
    test_txt_path = os.path.join(config.RAW_DATA_DIR, 'test_subject', 'test_simple.txt')
    with open(test_txt_path, 'w', encoding='utf-8') as f:
        f.write("Este es un archivo de texto plano.")
    print(f"Creado TXT de prueba: {test_txt_path}")

    # --- Ejemplos de uso del loader ---
    print("\n--- Probando DocumentLoader ---")

    test_pdf_path = None
    test_pdf_path = os.path.join(config.RAW_DATA_DIR, 'test_subject','resumen sobre trigonometría 12g.pdf')
    # Prueba con PDF nativo
    print(f"Probando PDF: {test_pdf_path}")
    if test_pdf_path and os.path.exists(test_pdf_path):
        pdf_text = loader.load_document(test_pdf_path)
        print(f"\nTexto de '{os.path.basename(test_pdf_path)}':\n{pdf_text[:100]}...")

    # Prueba con DOCX
    if 'test_docx_path' in locals() and os.path.exists(test_docx_path):
        docx_text = loader.load_document(test_docx_path)
        print(f"\nTexto de '{os.path.basename(test_docx_path)}':\n{docx_text[:100]}...")

    # Prueba con HTML
    if os.path.exists(test_html_path):
        html_text = loader.load_document(test_html_path)
        print(f"\nTexto de '{os.path.basename(test_html_path)}':\n{html_text[:100]}...")

    # Prueba con TXT
    if os.path.exists(test_txt_path):
        txt_text = loader.load_document(test_txt_path)
        print(f"\nTexto de '{os.path.basename(test_txt_path)}':\n{txt_text[:100]}...")

    # Ejemplo de un archivo que no existe
    non_existent_file = "/path/to/non_existent_file.pdf"
    loader.load_document(non_existent_file)